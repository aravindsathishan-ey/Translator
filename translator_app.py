import os
import uuid
import math
import time
from io import BytesIO
from PIL import Image
import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from openpyxl import load_workbook
from pptx import Presentation
from urllib.parse import quote, unquote
from datetime import datetime, timedelta, timezone
from azure.core.credentials import AzureKeyCredential, AzureNamedKeyCredential
from azure.data.tables import TableServiceClient, UpdateMode
from azure.ai.translation.document import DocumentTranslationClient, DocumentTranslationInput, TranslationTarget
from azure.storage.blob import BlobServiceClient, ContentSettings, generate_account_sas, ResourceTypes, AccountSasPermissions

load_dotenv()

ACCOUNT_NAME = os.getenv("AZURE_STORAGE_ACCOUNT_NAME")
ACCOUNT_KEY = os.getenv("AZURE_STORAGE_ACCOUNT_KEY")
SRC_CONTAINER = os.getenv("AZURE_SOURCE_CONTAINER", "transtest")
DST_CONTAINER = os.getenv("AZURE_TARGET_CONTAINER", "transtesttarget")
ENDPOINT = os.getenv("AZURE_DOCUMENT_TRANSLATION_ENDPOINT")
KEY = os.getenv("AZURE_DOCUMENT_TRANSLATION_KEY")
TABLE_NAME = os.getenv("AZURE_TABLE_NAME", "learningtesttable")

#DOC-CONFIG
A4_WIDTH_PT = 595
A4_HEIGHT_PT = 842
DEFAULT_MARGIN_PT = 72
DEFAULT_FONT_SIZE_PT = 12
LINE_HEIGHT_FACTOR = 1.2
AVG_CHARS_PER_LINE = 80          
PARA_SPACING_LINES = 0.3         
TABLE_ROW_LINES = 1    
TABLE_AFTER_SPACING_LINES = 0.5
HEADER_FOOTER_LINES = 1.0    
EMU_PER_POINT = 12700


class Translator:
    def __init__(self):
        self.sas_token = self.generate_sas()
        self.blob_service = BlobServiceClient(f"https://{ACCOUNT_NAME}.blob.core.windows.net", credential=ACCOUNT_KEY)
        self.src_client = self.blob_service.get_container_client(SRC_CONTAINER)
        self.dst_client = self.blob_service.get_container_client(DST_CONTAINER)
        self.SOURCE_URL = f"https://{ACCOUNT_NAME}.blob.core.windows.net/{SRC_CONTAINER}?{self.sas_token}"
        self.TARGET_URL = f"https://{ACCOUNT_NAME}.blob.core.windows.net/{DST_CONTAINER}?{self.sas_token}"
        self.translator = DocumentTranslationClient(ENDPOINT, AzureKeyCredential(KEY))
        self.table_client = self.table_init()
        self._wipe_container(self.src_client)
        self._wipe_container(self.dst_client)

    def _safe_blob_name(self, name: str) -> str:
        return quote(name, safe="~()*!.'-_")

    def _wipe_container(self, client):
        try:
            for blob in client.list_blobs():
                try:
                    client.delete_blob(blob.name)
                except Exception:
                    pass
        except Exception:
            pass

    #blob table init
    def table_init(self):
        table_client = None
        try:
            table_credential = AzureNamedKeyCredential(ACCOUNT_NAME, ACCOUNT_KEY)
            table_service = TableServiceClient(
                        endpoint=f"https://{ACCOUNT_NAME}.table.core.windows.net",
                        credential=table_credential
            )
            return table_service.create_table_if_not_exists(TABLE_NAME)
        except Exception:
            return table_client

    #generate ac levl sas token
    def generate_sas(self):
        try:
            start = datetime.now(timezone.utc) - timedelta(minutes=5)
            expiry = datetime.now(timezone.utc) + timedelta(hours=1)

            return generate_account_sas(
                account_name=ACCOUNT_NAME,
                account_key=ACCOUNT_KEY,
                resource_types=ResourceTypes(object=True, container=True, service=True),
                permission=AccountSasPermissions(
                    read=True, write=True, list=True, delete=True, create=True, add=True
                ),
                start=start,
                expiry=expiry,
            )
        except Exception as e:
            raise RuntimeError(f"Failed to generate SAS token: {e}")

    def upload_files(self, files):
        uploaded_blob_names = []
        for f in files:
            try:
                blob_name = f"{uuid.uuid4()}-{f.name}"
                safe_blob_name = self._safe_blob_name(blob_name)
                blob_client = self.src_client.get_blob_client(safe_blob_name)

                blob_client.upload_blob(
                    f,
                    overwrite=True,
                    content_settings=ContentSettings(content_type=f.type)
                )
                uploaded_blob_names.append(safe_blob_name)

            except Exception as e:
                st.error(f"Upload failed for {f.name}: {e}")

        return uploaded_blob_names

    def translate(self, target_lan="en"):
        try:
            return self.translator.begin_translation(
                self.SOURCE_URL,
                self.TARGET_URL,
                target_language= target_lan
            )
        except Exception as e:
            raise RuntimeError(f"Translation failed: {e}")
    
    def page_count(self, file_bytes, extension):
        try:
            # print("ccccccccccccccccccccc")
            ext = extension.lower()
            if ext == ".pdf":
                reader = PdfReader(BytesIO(file_bytes))
                return len(reader.pages)
            elif ext == ".pptx":
                prs = Presentation(BytesIO(file_bytes))
                return len(prs.slides)
            elif ext == ".docx":
                from docx import Document
                return self.estimate_docx_a4_pages(file_bytes)  
            elif ext == ".txt":
                return self.estimate_txt_a4_pages(file_bytes)
            elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"]:
                with Image.open(BytesIO(file_bytes)) as img:
                    return getattr(img, "n_frames", 1)
            elif ext == ".xlsx":
                # print("ggggggggggggggggg")
                return self.estimate_excel_a4_pages(file_bytes)
            else:
                return None
        except Exception as e:
            return None

    
    def estimate_docx_a4_pages(self, docx_bytes):
        """
        Estimate DOCX pages (works with images, tables, paragraphs, manual page breaks).
        Pure Python: uses height estimation rather than true layout.
        """
        doc = Document(BytesIO(docx_bytes))
        def _has_page_break(paragraph) -> bool:
            """
            Detects if a Word paragraph contains a manual page break:
            looks for <w:br w:type="page"/> anywhere in the paragraph XML.
            """
            p = paragraph._element
            br_elems = p.xpath('.//w:br[@w:type="page"]')
            return len(br_elems) > 0
        def _to_points_len(obj) -> float:
            """
            Convert python-docx length-like value to points (float) safely.
            - python-docx Length objects have .pt
            - Some values may already be numeric (EMU)
            """
            try:
                return float(obj.pt)
            except Exception:
                pass

            try:
                val = float(obj)
                if val > 10000:
                    return val / EMU_PER_POINT
                return val
            except Exception:
                # Fallback
                return 0.0

        try:
            section = doc.sections[0]
            page_height_pt = _to_points_len(section.page_height)
            top_margin_pt = _to_points_len(section.top_margin)
            bottom_margin_pt = _to_points_len(section.bottom_margin)
            printable_height = max(1.0, page_height_pt - (top_margin_pt + bottom_margin_pt))
        except Exception:
            printable_height = A4_HEIGHT_PT - (2 * DEFAULT_MARGIN_PT)

        line_height = DEFAULT_FONT_SIZE_PT * LINE_HEIGHT_FACTOR
        total_height_pt = 0.0
        explicit_page_breaks = 0

        for para in doc.paragraphs:
            #Count manual page breaks in this paragraph
            if _has_page_break(para):
                explicit_page_breaks += 1

            text = (para.text or "").strip()

            if not text:
                total_height_pt += line_height
                continue

            char_count = len(text)
            wrapped_lines = max(1, math.ceil(char_count / AVG_CHARS_PER_LINE))
            total_height_pt += wrapped_lines * line_height

            # Add a bit of after-paragraph spacing
            total_height_pt += PARA_SPACING_LINES * line_height

        # 2) TABLES (count rows)
        for table in doc.tables:
            row_count = len(table.rows)
            if row_count > 0:
                total_height_pt += row_count * (TABLE_ROW_LINES * line_height)
                total_height_pt += TABLE_AFTER_SPACING_LINES * line_height

        # 3) IMAGES (inline shapes)
        for shape in getattr(doc, "inline_shapes", []):
            # Height in points
            h_pt = _to_points_len(shape.height)
            if h_pt <= 0:
                # conservative fallback if height cannot be read
                h_pt = 150.0
            total_height_pt += h_pt + (0.5 * line_height)

        try:
            header = section.header
            footer = section.footer
            if header and header.paragraphs and any((p.text or "").strip() for p in header.paragraphs):
                total_height_pt += HEADER_FOOTER_LINES * line_height
            if footer and footer.paragraphs and any((p.text or "").strip() for p in footer.paragraphs):
                total_height_pt += HEADER_FOOTER_LINES * line_height
        except Exception:
            pass

        # 5) Convert accumulated height to pages
        content_pages = math.ceil(max(1.0, total_height_pt) / printable_height)

        # Add explicit manual page breaks (each forces a new page)
        pages = content_pages + explicit_page_breaks

        # Safety bound
        return max(1, int(pages))

    def estimate_excel_a4_pages(self, xl_bytes):
        wb = load_workbook(BytesIO(xl_bytes), data_only=True)
        printable_width = A4_WIDTH_PT - (2 * DEFAULT_MARGIN_PT)
        printable_height = A4_HEIGHT_PT - (2 * DEFAULT_MARGIN_PT)

        total_pages = 0

        for sheet in wb.worksheets:
            #detect actual used cell range
            min_row = None
            max_row = None
            min_col = None
            max_col = None

            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value not in (None, ""):
                        r = cell.row
                        c = cell.column

                        min_row = r if min_row is None else min(min_row, r)
                        max_row = r if max_row is None else max(max_row, r)
                        min_col = c if min_col is None else min(min_col, c)
                        max_col = c if max_col is None else max(max_col, c)

            #If sheet empty → 1 page
            if min_row is None:
                total_pages += 1
                continue

            used_cols = max_col - min_col + 1
            used_rows = max_row - min_row + 1

            # Estimate dimensions
            total_width = used_cols * 64      # approx 64pt per column
            total_height = used_rows * 15     # approx 15pt per row

            width_pages = math.ceil(total_width / printable_width)
            height_pages = math.ceil(total_height / printable_height)

            total_pages += max(1, width_pages * height_pages)

        return total_pages
    
    def estimate_txt_a4_pages(self, file_bytes):
        text = file_bytes.decode("utf-8", errors="ignore")
        total_chars = len(text)

        #printable area
        printable_width = A4_WIDTH_PT - 2 * DEFAULT_MARGIN_PT
        printable_height = A4_HEIGHT_PT - 2 * DEFAULT_MARGIN_PT

        # Typography estimation
        char_width = DEFAULT_FONT_SIZE_PT * 0.5
        line_height = DEFAULT_FONT_SIZE_PT * LINE_HEIGHT_FACTOR

        chars_per_line = printable_width / char_width
        lines_per_page = printable_height / line_height

        chars_per_page = chars_per_line * lines_per_page

        estimated_pages = math.ceil(total_chars / chars_per_page)

        return max(1, estimated_pages)

    def download_translated(self):
        translated_files = []
        for blob in self.dst_client.list_blobs():
            file_bytes = self.dst_client.get_blob_client(blob.name).download_blob().readall()

            decoded = unquote(blob.name)
            cleaned = "-".join(decoded.split("-")[1:])

            if "." in cleaned:
                base, ext = cleaned.rsplit(".", 1)
                final_name = f"{target_language}_{base}_translated.{ext}"
            else:
                final_name = f"{target_language}_{cleaned}_translated"

            translated_files.append((final_name, file_bytes))
            # translated_files.append((meta, file_bytes))
        return translated_files

    #blob lvl cleanup
    def cleanup(self, uploaded):
        #src
        for name in uploaded:
            try:
                self.src_client.delete_blob(name)
            except Exception:
                pass

        #dst
        for blob in self.dst_client.list_blobs():
            try:
                self.dst_client.delete_blob(blob.name)
            except Exception:
                pass


client = Translator()



# --- LOGOS ---

st.set_page_config(
    page_title="Doc-Translator",
    page_icon="🌐",
    layout="centered",
    initial_sidebar_state="collapsed"
)

st.markdown(
    """
    <style>
    /* Main app background */
    .stApp {
        background-color: #000000;
    }

    /* Top header bar */
    header[data-testid="stHeader"] {
        background-color: #000000;
    }

    /* Toolbar (menu dots area) */
    div[data-testid="stToolbar"] {
        background-color: #000000;
    }

    /* Optional: remove header bottom border */
    header[data-testid="stHeader"]::after {
        background: none;
    }

    /* Text color safety */
    body {
        color: white;
    }
    </style>
    """,
    unsafe_allow_html=True
)




LOGO_URL = "https://th.bing.com/th/id/OIP.Wgr313SqtaL6NaKsDJfihQAAAA?o=7rm=3&rs=1&pid=ImgDetMain&o=7&rm=3"
# One centered row
left, mid, right = st.columns([1, 3, 1])
with mid:
    st.markdown(
        f"""
        <div style="
            display: flex;
            align-items: flex-end;
            gap: 16px;
            margin-bottom: 24px;
        ">
            <img src="{LOGO_URL}" style="height:72px;">
            <h1 style="
                margin: 0;
                padding: 0;
                font-size: 42px;
                font-weight: 600;
                line-height: 1;
                color: white;
            ">
                Doc-Translator
            </h1>
        </div>
        """,
        unsafe_allow_html=True
    )


if "translated_files" not in st.session_state:
    # Will store: list[tuple[str, bytes]]
    st.session_state.translated_files = []
if "translate_stats" not in st.session_state:
    # Will store success/fail counts
    st.session_state.translate_stats = None
if "uploaded_temp_refs" not in st.session_state:
    # Keep uploaded_names for cleanup later
    st.session_state.uploaded_temp_refs = None

# --- UI controls ---
uploaded_files = st.file_uploader(
    "Upload files",
    type=["pdf", "docx", "txt", "html", "xlsx"],
    accept_multiple_files=True,
    key="uploader"  # stable key helps retain value across reruns
)

target_language = st.text_input("Target language (ISO code)", value="en")

rail_left, rail_right = st.columns([1, 1])
with rail_left:
    # -------- Buttons row inside the SAME rail --------
    # Two equal columns: left = Translate, right = Clear
    c_left, c_right = st.columns(2)

    with c_left:
        translate_clicked = st.button("Translate", key="translate_btn")
with rail_right:
    d_left, d_right = st.columns([4,1])

    with d_right:
        # This button will naturally align to the right edge of the rail
        clear_clicked = st.button("Clear", key="clear_btn")

# ---------------- EY Yellow styles ----------------
st.markdown("""
<style>
/* Make all st.button elements look like EY buttons and fill their column */
.stButton > button {
    width: 100% !important;                 /* fill column width */
    height: 44px !important;                /* consistent height */
    background-color: #FFCF00 !important;   /* EY Yellow */
    color: #000000 !important;              /* black text */
    border: 1px solid #E6B800 !important;   /* EY gold */
    border-radius: 10px !important;
    font-weight: 700 !important;
}
.stButton > button:hover {
    background-color: #FFD84D !important;   /* lighter EY yellow on hover */
    color: #000000 !important;
}
</style>
""", unsafe_allow_html=True)
#
# --- Clear behavior: wipe state AND cleanup temps (once) ---
if clear_clicked:
    try:
        if st.session_state.uploaded_temp_refs:
            client.cleanup(st.session_state.uploaded_temp_refs)
    except Exception as e:
        st.info(f"Cleanup note: {e}")
    st.session_state.translated_files = []
    st.session_state.translate_stats = None
    st.session_state.uploaded_temp_refs = None
    # Also reset the uploader widget by changing its key (optional trick):
    # st.session_state.uploader = None
    st.rerun()

# --- Run translation only when pressed ---
if translate_clicked and uploaded_files:
    with st.spinner("Uploading files..."):
        uploaded_names = client.upload_files(uploaded_files)

    #####table insertion
    row_keys = []

    for f, blob_name in zip(uploaded_files, uploaded_names):
        f.seek(0)
        file_bytes = f.read()
        filename = f.name.lower()
        extension = "." + filename.split(".")[-1]
        # print("lsdsdsdd",file_bytes)
        page_count = client.page_count(file_bytes, extension)
        # print("page______count", page_count, extension)
        row_key = str(uuid.uuid4())
        row_keys.append(row_key)

        if client.table_client:
            client.table_client.upsert_entity({
                "PartitionKey": "files",
                "RowKey": row_key,
                "original_name": f.name,
                "blob_name": blob_name,
                "file_type": filename.split(".")[-1],
                "target_language": target_language,
                "page_count": page_count,
                "status": "Uploaded",
                "uploaded_on": datetime.now(timezone.utc).isoformat(),
                "translated_on": ""
            })

    # st.success(f"Uploaded {len(uploaded_names)} file(s).")

    poller = client.translate(target_language)
    progress_bar = st.progress(0)
    status_placeholder = st.empty()

    with st.spinner("Translating..."):
        while not poller.done():

            try:
                details = poller.details

                if details and details.summary:
                    total = details.summary.get("total", 0)
                    succeeded = details.summary.get("succeeded", 0)
                    failed = details.summary.get("failed", 0)
                    completed = succeeded + failed

                    if total > 0:
                        progress = int((completed / total) * 100)
                        progress_bar.progress(progress)

                    status_placeholder.info(
                        f"Processed {completed} of {total} documents..."
                    )
                else:
                    status_placeholder.info("Initializing translation job...")

            except:
                status_placeholder.info("Preparing documents...")

            time.sleep(2)
        poller.result()

    progress_bar.progress(100)
    status_placeholder.success("Translation completed successfully!")

    for row_key in row_keys:
        client.table_client.update_entity({
            "PartitionKey": "files",
            "RowKey": row_key,
            "status": "Translated",
            "translated_on": datetime.now(timezone.utc).isoformat()
        }, mode=UpdateMode.MERGE)

    # Save stats in state so they persist across reruns
    st.session_state.translate_stats = {
        "succeeded": getattr(poller.details, "documents_succeeded_count", None),
        "failed": getattr(poller.details, "documents_failed_count", None),
    }

    # Fetch translated payloads and persist in state
    translated = list(client.download_translated())  # ensure it's materialized
    # Each item should be (fname, bytes)
    st.session_state.translated_files = translated

    # Save uploaded temp refs for later explicit cleanup on 'Clear'
    st.session_state.uploaded_temp_refs = uploaded_names

# --- Show results if any are in session_state ---
if st.session_state.translate_stats:
    st.write(f"Documents succeeded: {st.session_state.translate_stats['succeeded']}")
    st.write(f"Documents failed: {st.session_state.translate_stats['failed']}")

if st.session_state.translated_files:
    st.subheader("Download Translated Files")
    for i, (fname, data) in enumerate(st.session_state.translated_files):
        # give every download button a unique, stable key
        st.download_button(
            label=f"Download {fname}",
            data=data,
            file_name=fname,
            key=f"dl_{i}_{fname}"
        )

    st.info("Use **Clear** to remove temporary files and reset the page.")
else:
    st.caption("No translated files yet. Upload and click **Translate**.")
